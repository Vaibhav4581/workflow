from docx import Document
from pymongo import MongoClient
import bcrypt
import re

mongo_url = "mongodb://Adisankar:CB1E9r7mjPV5YLpq@ac-irtrr2z-shard-00-00.3nnx8jj.mongodb.net:27017,ac-irtrr2z-shard-00-01.3nnx8jj.mongodb.net:27017,ac-irtrr2z-shard-00-02.3nnx8jj.mongodb.net:27017/?ssl=true&replicaSet=atlas-czuifj-shard-0&authSource=admin&appName=Cluster0"

client = MongoClient(mongo_url)
db = client.get_database("test")
users_collection = db.users

db_users = list(users_collection.find({}, {"email": 1, "name": 1}))
db_emails = {u.get("email", "").strip().lower() for u in db_users if u.get("email")}
db_names = {u.get("name", "").strip().lower() for u in db_users if u.get("name")}

doc = Document('SNGCE Staff Directory.docx')

# Extract departments from tables
# But tables are not directly linked to paragraphs. We can just use the table index.
# In inspect_docx, we saw:
# Table 0: Principal
# Table 1: Civil
# Let's just find the preceding paragraph.
# Docx structure: paragraphs and tables are interleaved.
# A better way is to iterate block level elements. But python-docx doesn't easily interleave.
# I'll just assign a default department or infer from doc.paragraphs.

departments = []
for p in doc.paragraphs:
    if "department of" in p.text.lower():
        # clean it
        dept = p.text.replace("Department of", "").strip()
        departments.append(dept)

# It's easier to just use standard abbreviations or the full string.
# But let's just insert them as "Faculty" with department "Unknown" and they can update it,
# OR I can just map them if I know.
# Wait, for now I will just use 'Staff' or 'Faculty' role.
# Users missing are specific. Let's just create them.

missing_people = []
seen_emails = set()

# Mapping table index to department is hard without interleaving, so let's just use "Faculty" as dept or "Various".
default_password = "password123".encode('utf-8')
hashed = bcrypt.hashpw(default_password, bcrypt.gensalt()).decode('utf-8')

import datetime

def generate_email(name):
    # e.g. "Mr. John Doe" -> "johndoe@sngce.ac.in"
    clean = re.sub(r'^(Mr\.|Ms\.|Dr\.|Prof\.)\s*', '', name, flags=re.IGNORECASE)
    clean = re.sub(r'[^a-zA-Z]', '', clean).lower()
    return f"{clean}@sngce.ac.in"

docs_to_insert = []

for table in doc.tables:
    header_row = table.rows[0]
    email_idx = -1
    name_idx = -1
    for idx, cell in enumerate(header_row.cells):
        text = cell.text.strip().lower()
        if "email" in text:
            email_idx = idx
        if "name" in text and "faculty" not in text:
            name_idx = idx
            
    if name_idx == -1:
        for idx, cell in enumerate(header_row.cells):
            text = cell.text.strip().lower()
            if "name" in text:
                name_idx = idx
                
    if email_idx == -1 or name_idx == -1:
        continue
        
    for row in table.rows[1:]:
        cells = row.cells
        if len(cells) > max(email_idx, name_idx):
            name = cells[name_idx].text.replace('\n', '').strip()
            email = cells[email_idx].text.replace('\n', '').strip().lower()
            
            if not name and not email:
                continue
                
            if email in seen_emails:
                continue
            if email:
                seen_emails.add(email)
                
            found = False
            if email and email in db_emails:
                found = True
            elif name and name.lower() in db_names:
                found = True
                
            if not found:
                # Add to DB
                # Split name into fName and lName
                clean_name = re.sub(r'^(Mr\.|Ms\.|Dr\.|Prof\.)\s*', '', name, flags=re.IGNORECASE).strip()
                parts = clean_name.split(' ', 1)
                fName = parts[0] if len(parts) > 0 else clean_name
                lName = parts[1] if len(parts) > 1 else ""
                
                final_email = email if email else generate_email(name)
                
                # Check again to avoid dup generated emails
                if final_email in db_emails:
                    final_email = f"user{len(docs_to_insert)}_" + final_email
                
                user_doc = {
                    "fName": fName,
                    "lName": lName,
                    "name": name, # Some places use 'name'
                    "email": final_email,
                    "password": hashed,
                    "role": "Faculty",
                    "department": "Other", # default
                    "createdAt": datetime.datetime.utcnow(),
                    "updatedAt": datetime.datetime.utcnow(),
                    "__v": 0
                }
                docs_to_insert.append(user_doc)
                db_emails.add(final_email) # prevent dups

if docs_to_insert:
    res = users_collection.insert_many(docs_to_insert)
    print(f"Inserted {len(res.inserted_ids)} users successfully!")
else:
    print("No missing users to insert.")
