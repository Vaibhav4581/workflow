from docx import Document
import sys

try:
    doc = Document('SNGCE Staff Directory.docx')
    for i, para in enumerate(doc.paragraphs[:50]):
        if para.text.strip():
            print(f"Para {i}: {para.text.strip()}")
            
    print("\nTables:")
    for t_idx, table in enumerate(doc.tables[:2]):
        print(f"\nTable {t_idx}:")
        for r_idx, row in enumerate(table.rows[:5]):
            cells = [cell.text.strip() for cell in row.cells]
            print(f"Row {r_idx}: {cells}")
except Exception as e:
    print(f"Error: {e}")
